

def basic_calculator(student_numbers):
    """
    Calculate the sum of student numbers.

    Args:
    student_numbers (list of int): The list of student numbers.

    Returns:
    Int: The sum of student numbers.
    """
    # Calculate the sum of the student numbers
    sum_of_numbers = sum(student_numbers)

    return sum_of_numbers

def extract_course_marks(student_numbers):
    '''
    Extract course marks to be used for grading
    Args: 
    Returns:
    list: The list of course marks extracted from the sum of student numbers.
    '''
    student_numbers_sum = basic_calculator(student_numbers)
     # Extract course marks from the sum (ignoring the first digit)
    str_sum = str(student_numbers_sum)[1:]  # Convert to string and ignore first digit
    course_marks = [int(str_sum[i:i+2]) for i in range(0, len(str_sum), 2) if i+2 <= len(str_sum)]

    return course_marks


# Update the map_marks_to_grades function to use the provided grading scale
def map_marks_to_grades(marks, grading_scale):
    """
    Map the course marks to grades based on the provided grading scale.

    Args:
    marks (list of int): The list of course marks.
    grading_scale (dict): The grading scale mapping mark ranges to letter grades.

    Returns:
    list: The list of grades corresponding to the marks.
    """
    grades = []
    for mark in marks:
        for mark_range, grade in grading_scale.items():
            if mark in mark_range:
                grades.append(grade)
                break
        else:
            # If no grade is found, assume it's the lowest grade
            grades.append('F')
    
    return grades


# the calculate_cgpa function 
def calculate_cgpa(course_marks, grades, course_credit_units):
    """
    Calculate the CGPA based on course marks, grade points, and course credit units.

    Args:
    course_marks (dict): A dictionary of course codes and their respective marks.
    grade_points (dict): The mapping of letter grades to grade points.
    course_credit_units (dict): The credit units for each course code.

    Returns:
    float: The calculated CGPA.
    """
    # Calculate the total weighted grade points and the total credit units
    total_weighted_grade_points = sum(grade_points[grade] * course_credit_units[course_code]
                                      for course_code, grade in zip(course_marks, grades))
    total_units = sum(course_credit_units.values())

    # Compute the CGPA
    cgpa = total_weighted_grade_points / total_units if total_units > 0 else 0
    return cgpa


def save_results_to_file(computed_cgpa, course_codes, generated_course_marks, grading_scale, student_numbers, grade_points, credit_units):
    from docx import Document
    import os

    doc_path = f"{os.getcwd()}\CGPA_Calculation_Results.docx"

    if os.path.exists(doc_path):
        doc = Document(doc_path)
        doc.add_heading('a) CGPA Calculation Results', 0)
        make_word_document_data(doc, computed_cgpa, course_codes, generated_course_marks, grading_scale, student_numbers, grade_points, credit_units)
        
    else:
        # Create a new Word document
        doc = Document()
        doc.add_heading('a) CGPA Calculation Results', 0)
        make_word_document_data(doc, computed_cgpa, course_codes, generated_course_marks, grading_scale, student_numbers, grade_points, credit_units)
       
    # Save the document
    doc.save(doc_path)

def make_word_document_data(doc, 
                            computed_cgpa, 
                            course_codes, 
                            generated_course_marks, 
                            grading_scale,
                            student_numbers,
                            grade_points,
                            credit_units
                        ):

    # Add the CGPA to the document
    doc.add_heading(f'Computed CGPA: {computed_cgpa}', level=1)
    doc.add_heading(f"Students' Numbers: {student_numbers}", level=3)
    doc.add_heading(f"Sum of Students' Numbers: {basic_calculator(student_numbers)}", level=3)

    # Add the course marks and their corresponding grades to the document
    doc.add_heading('Course Marks and Grades Usung Functions', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Course Code'
    hdr_cells[1].text = 'Credit Unit'
    hdr_cells[2].text = 'Mark'
    hdr_cells[3].text = 'Grade'
    hdr_cells[4].text = 'Grade Score'


    # Add a row for each course
    for course_code, mark in zip(course_codes, generated_course_marks):
        grade = map_marks_to_grades([mark], grading_scale)[0]  # Get the grade for the current mark
        print(course_code, mark, credit_units[course_code], grade_points[grade])
        row_cells = table.add_row().cells
        row_cells[0].text = course_code
        row_cells[1].text = str(credit_units[course_code])
        row_cells[2].text = str(mark)
        row_cells[3].text = grade
        row_cells[4].text = str(grade_points[grade])

    # Add an empty paragraph
    doc.add_paragraph()


# STATIC DATA
# Courses | We shall use Course Codes
course_codes = ['CSK 1101', 'CSC 1102', 'CSC 1104', 'CSC 1105']

# Student Numbers for this group
student_numbers = [1900717625, 2300717623, 2300707764, 2300713410, 207006808]

# Provided grading scale
grading_scale = {
    range(90, 101): 'A+', range(80, 90): 'A', range(75, 80): 'B+', range(70, 75): 'B',
    range(65, 70): 'C+', range(60, 65): 'C', range(55, 60): 'D+', range(50, 55): 'D',
    range(45, 50): 'E', range(40, 45): 'E-', range(0, 40): 'F'
}

# Grade points corresponding to the letter grades
grade_points = {
    'A+': 5.0, 'A': 5.0, 'B+': 4.5, 'B': 4.0, 'C+': 3.5, 'C': 3.0, 'D+': 2.5, 'D': 2.0, 
    'E': 1.5, 'E-': 1.0, 'F': 0.0
}

# Provided credit units for the courses
credit_units = {
    'CSK 1101': 4, 'CSC 1102': 4, 'CSC 1104': 4, 'CSC 1105': 4
}


# Calculate the CGPA using the mapped grades and the provided credit units

sum_of_numbers = basic_calculator(student_numbers)
print('Sum of Numbers: ', sum_of_numbers)

# Generate marks from the student numbers
generated_course_marks = extract_course_marks(student_numbers)
print('Course Marks: ', generated_course_marks)

# Map the generated marks to the course codes
course_marks = dict(zip(course_codes, generated_course_marks))
print('course_marks: ', course_marks)

# Map the generated marks to grades
grades = map_marks_to_grades(generated_course_marks, grading_scale)
print('Grades: ', grades)

grade_mapping = dict(zip(course_codes, grades))
print('Course Grades : ', grade_mapping)

# Calculate CGPA
cgpa = calculate_cgpa(course_marks, grades, credit_units)

print()
print('CGPA for individual fuctions: ', cgpa)
print()


# Save results to file
save_results_to_file(cgpa, course_codes, generated_course_marks, grading_scale, student_numbers, grade_points, credit_units)
