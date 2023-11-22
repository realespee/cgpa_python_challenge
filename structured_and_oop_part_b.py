
# Defining the base class for a generic BasicCalculator
# Base calculator class
class BasicCalculator:
    """
    A generic calculator class with basic arithmetic operations.
    """
    def add(self, a, b):
        return a + b

    def subtract(self, a, b):
        return a - b

    def multiply(self, a, b):
        return a * b

    def divide(self, a, b):
        if b == 0:
            raise ValueError("Cannot divide by zero")
        return a / b

# Define the CGPACalculator class that inherits from Calculator
class CGPACalculator(BasicCalculator):
    """
    A CGPA calculator that extends the basic Calculator functionalities.
    """
    def __init__(self, grade_scale, grade_points, credit_units, student_numbers):
        super().__init__()
        self.grade_scale = grade_scale
        self.grade_points = grade_points
        self.credit_units = credit_units
        self.student_numbers = student_numbers
        self.course_codes = list(self.credit_units.keys())  # Assuming the order of courses matches the order of grades


    def sum_student_numbers(self, student_numbers):
        sum = 0
        for num in student_numbers:
            sum = self.add(sum, num)
        return sum
    
    def generate_course_marks(self, student_numbers):
        student_numbers_sum = self.sum_student_numbers(student_numbers)
        # Extract course marks from the sum (ignoring the first digit)
        str_sum = str(student_numbers_sum)[1:]  # Convert to string and ignore first digit
        course_marks = [int(str_sum[i:i+2]) for i in range(0, len(str_sum), 2) if i+2 <= len(str_sum)]
        return course_marks
    
    def map_marks_to_grades(self, marks):
        """
        Map the course marks to grades based on the grading scale.
        """
        grades = []
        for mark in marks:
            for mark_range, grade in self.grade_scale.items():
                if mark in mark_range:
                    grades.append(grade)
                    break
            else:
                grades.append('F')
        return grades
    
    def make_word_document_data(self, doc):
        computed_cgpa = self.calculate_cgpa()
        # Add the CGPA to the document
        doc.add_heading(f'CGPA: {computed_cgpa}', level=1)
        doc.add_heading(f"Students' Numbers: {self.student_numbers}", level=3)
        doc.add_heading(f"Sum of Students' Numbers: {self.sum_student_numbers(self.student_numbers)}", level=3)

        # Add the course marks and their corresponding grades to the document
        doc.add_heading('Course Marks and Grades Usung Classes', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Add the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Course Code'
        hdr_cells[1].text = 'Credit Unit'
        hdr_cells[2].text = 'Mark'
        hdr_cells[3].text = 'Grade'
        hdr_cells[4].text = 'Grade Score'

        # Add a row for each course
        for course_code, mark in zip(self.course_codes, self.generate_course_marks(self.student_numbers)):
            credit_units = self.credit_units
            grade_points = self.grade_points
            grade = self.map_marks_to_grades([mark])[0]  # Get the grade for the current mark
            row_cells = table.add_row().cells
            row_cells[0].text = course_code
            row_cells[1].text = str(credit_units[course_code])
            row_cells[2].text = str(mark)
            row_cells[3].text = grade
            row_cells[4].text = str(grade_points[grade])

        # Add an empty paragraph
        doc.add_paragraph()
    
    def save_results_to_file(self):
        from docx import Document
        import os

        doc_path = f"{os.getcwd()}\CGPA_Calculation_Results.docx"

        if os.path.exists(doc_path):
            doc = Document(doc_path)
            doc.add_heading('b) CGPA Calculation Results', 0)
            self.make_word_document_data(doc)
        else:
            # Create a new Word document
            doc = Document()
            doc.add_heading('b) CGPA Calculation Results', 0)
            self.make_word_document_data(doc)
        
        # Save the document
        doc.save(doc_path)

    def calculate_cgpa(self):
        """
        Calculate the CGPA based on course marks.
        """
        marks = self.generate_course_marks(self.student_numbers)
        grades = self.map_marks_to_grades(marks)
        total_weighted_grade_points = 0
        total_units = 0

        for course_code, grade in zip(self.course_codes, grades):
            grade_point = self.grade_points[grade]
            credit_unit = self.credit_units[course_code]
            total_weighted_grade_points = self.add(
                total_weighted_grade_points,
                self.multiply(grade_point, credit_unit)
            )
            total_units = self.add(total_units, credit_unit)

        cgpa = self.divide(total_weighted_grade_points, total_units)
        return cgpa




# Define the grading scale, grade points, credit units, and students' numbers
grading_scale = {
    range(90, 101): 'A+', range(80, 90): 'A', range(75, 80): 'B+', range(70, 75): 'B',
    range(65, 70): 'C+', range(60, 65): 'C', range(55, 60): 'D+', range(50, 55): 'D',
    range(45, 50): 'E', range(40, 45): 'E-', range(0, 40): 'F'
}

grade_points = {
    'A+': 5.0, 'A': 5.0, 'B+': 4.5, 'B': 4.0, 'C+': 3.5, 'C': 3.0, 'D+': 2.5, 'D': 2.0, 
    'E': 1.5, 'E-': 1.0, 'F': 0.0
}

credit_units = {
    'CSK 1101': 4, 'CSC 1102': 4, 'CSC 1104': 4, 'CSC 1105': 4
}

student_numbers = [1900717625, 2300717623, 2300707764, 2300713410, 207006808]

# Instantiate the CGPA calculator
cgpa_calculator = CGPACalculator(grading_scale, grade_points, credit_units, student_numbers)

# Calculate the CGPA using the refactored code with inheritance
computed_cgpa = cgpa_calculator.calculate_cgpa()
print()
print('CGPA: ', computed_cgpa)
print()

#  Save Results to File
cgpa_calculator.save_results_to_file()